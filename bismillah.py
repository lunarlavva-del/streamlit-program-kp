import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import win32com.client
import pythoncom  # Import pythoncom untuk CoInitialize
import os

# Fungsi untuk mengganti placeholder dalam template Word
def fill_template(template_path, output_docx_path, data):
    doc = Document(template_path)  # Pastikan path template benar

    # Flag untuk mengetahui apakah kita sedang di bagian tanda tangan
    tanda_tangan_section = False

    # Ganti placeholder di paragraf
    for para in doc.paragraphs:
        if "Demikian berita acara ini kami buat" in para.text:  # Deteksi awal tanda tangan
            tanda_tangan_section = True

        for key, value in data.items():
            if key in para.text:  # Jika ditemukan placeholder
                para.text = para.text.replace(key, value)  # Ganti dengan data

                # Tambahkan format bold dan underline hanya di bagian tanda tangan
                if key == "nama_ttd" and tanda_tangan_section:
                    run = para.runs[0]  # Ambil bagian teks yang berisi placeholder
                    run.font.bold = True  # Set bold
                    run.font.underline = True  # Set underline

    # Ganti placeholder di tabel (jika ada)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if key in cell.text:  # Jika ditemukan placeholder di tabel
                        cell.text = cell.text.replace(key, value)  # Ganti dengan data

                        # Tambahkan format bold dan underline hanya di bagian tanda tangan
                        if key == "nama_ttd" and tanda_tangan_section:
                            run = cell.paragraphs[0].runs[0]  # Ambil bagian teks yang berisi placeholder
                            run.font.bold = True  # Set bold
                            run.font.underline = True  # Set underline

    # Simpan hasil perubahan ke file baru
    doc.save(output_docx_path)  # Menyimpan dokumen yang telah diubah ke path baru
    return output_docx_path

# Fungsi untuk mengonversi Word ke PDF menggunakan Microsoft Word
def convert_to_pdf(input_docx_path, output_pdf_path):
    pythoncom.CoInitialize()  # Memanggil CoInitialize sebelum menggunakan win32com
    
    # Verifikasi apakah file Word ada
    if not os.path.exists(input_docx_path):
        raise FileNotFoundError(f"File {input_docx_path} tidak ditemukan.")
    
    # Menggunakan Microsoft Word untuk konversi ke PDF
    word = win32com.client.Dispatch("Word.Application")
    doc = word.Documents.Open(input_docx_path)
    doc.SaveAs(output_pdf_path, FileFormat=17)  # 17 adalah format PDF di Word
    doc.Close()
    word.Quit()
    return output_pdf_path

# Streamlit UI
st.title("Generate Berita Acara (PDF)")
st.write("Isi form di bawah untuk menghasilkan berita acara dalam format PDF:")

# Form input
with st.form("form"):
    nama_ttd = st.text_input("Nama Penandatangan")
    jabatan_ttd = st.text_input("Jabatan Penandatangan")
    instansi_ttd = st.text_input("Instansi Penandatangan")
    nama_dosen = st.text_input("Nama Dosen")
    nama_kegiatan = st.text_input("Nama Kegiatan")
    hari_tanggal = st.text_input("Hari, Tanggal (contoh: Selasa, 10 Januari 2025)")
    pukul = st.text_input("Pukul (contoh: 09:00 - 12:00)")
    tempat_kegiatan = st.text_input("Tempat Kegiatan")
    kota_ttd = st.text_input("Kota penanda-tanganan dokumen")
    tanggal_ttd = st.text_input("Tanggal penanda-tanganan dokumen")
    submit = st.form_submit_button("Generate PDF")

# Jika form disubmit
if submit:
    # Data untuk menggantikan placeholder
    data = {
        "nama_ttd": nama_ttd,
        "jabatan_ttd": jabatan_ttd,
        "instansi_ttd": instansi_ttd,
        "nama_dosen": nama_dosen,
        "nama_kegiatan": nama_kegiatan,
        "hari_tanggal": hari_tanggal,
        "pukul": pukul,
        "tempat_kegiatan": tempat_kegiatan,
        "kota_ttd": kota_ttd,
        "tanggal_ttd": tanggal_ttd
    }

    # Path file template, hasil Word, dan hasil PDF
    template_path = r"D:\KP\BAP-Abdimas-Nareks-template_2.docx"  # Ganti dengan path template yang benar
    output_docx_path = r"D:\KP\berita_acara_filled.docx"  # Path output file Word yang sudah terisi
    output_pdf_path = r"D:\KP\berita_acara_filled.pdf"  # Path output file PDF

    # Isi template Word
    filled_docx = fill_template(template_path, output_docx_path, data)

    # Verifikasi file hasil pengisian template ada
    if os.path.exists(filled_docx):
        # Konversi ke PDF
        filled_pdf = convert_to_pdf(filled_docx, output_pdf_path)

        # Tombol unduh untuk PDF
        with open(filled_pdf, "rb") as file:
            st.download_button(
                label="Download Berita Acara (PDF)",
                data=file,
                file_name="berita_acara_filled.pdf",
                mime="application/pdf"
            )
    else:
        st.error(f"File Word {output_docx_path} tidak berhasil disimpan. Periksa kembali path atau template.")
